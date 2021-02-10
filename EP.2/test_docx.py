# test_docx.py
from docx import Document
import wikipedia

def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)

    # summary สำหรับบทความที่สรุป
    data = wikipedia.summary(keyword)

    # page + content บทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content


    doc = Document() #สร้างไฟล์ word ใน python
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2)
    doc.save(keyword + '.docx')
    print('สร้างไฟล์สำเร็จ')

try:
    Wiki('asdfasf','en')
except:
    print('ERROR')
    
#Wiki('ประเทศไทย','zh')
#Wiki('united state of america','en')
#Wiki('ปรระเทศญี่ปุ่น','jp')    

